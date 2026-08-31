# -*- coding: utf-8 -*-
"""Integrate the locally archived supplements for R0855, R0856 and R1040.

WHAT THIS IS FOR. `text_table_coverage.csv` records, for each of these three
records, that the supplement is archived and INSPECTED but that named parts of
it were never carried into the extraction tables. This reads those parts out of
the archived files and writes them as ordinary extraction rows.

WHAT IT REFUSES TO DO. Every value here is parsed out of the supplement file,
never retyped from a screen. Where the publisher suppressed a cell - R1040's
eTable 4 prints '*' for every stratum's case count and person-years - the row
is written WITHOUT a number and carries the reason, because a plausible
reconstruction of a suppressed count is worse than an absent one. Where the
supplement does not state something the schema wants - none of R0856's
supplementary tables name the reference category their hazard ratios are
measured against - the field says so in words rather than inheriting a guess
from the main paper.

WHAT IT MUST NOT TOUCH. `human_confirmed`, reviewer initials, consensus
`Resolved`, and `pool_eligible` belong to a person. Every row this writes is
human_confirmed=no, pool_eligible=no, human_gate=HUMAN_DUAL_EXTRACTION_REQUIRED.
Nothing here makes a record poolable; it only puts on the table what the
supplement says, so that a reviewer has something to confirm.
"""
import csv, hashlib, io, json, os, re, sys, datetime

OUT = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, OUT)
import bundle_paths
BASE = bundle_paths.BASE
SUPP = os.path.join(BASE, 'source_supplements')

#: PUBLISHER-DERIVED INPUTS LIVE OUTSIDE THIS FILE. The supplement filenames,
#: the covariate sentences quoted from each paper's footnote and the stratum
#: headings printed in its tables are the publications' content, not this
#: pipeline's; they sit in study_inputs.json, which is not published. What
#: stays here is the machinery: how a table is walked, what makes a cell
#: unreadable, and what the refusal says when it is.
INPUTS = bundle_paths.study_inputs()['supplements']
#: Per-record constants - study design, population, outcome, comparator -
#: which describe the PUBLICATIONS, not this pipeline. Data, like the
#: values themselves.
TPL = INPUTS['row_templates']
R1040 = INPUTS['r1040']
LBL = INPUTS['labels']
EXPO = {r: t['_exposure_classes'] for r, t in TPL.items()}
#: R0856's two supplementary tables, in the order they appear in the file.
R0856_CLASSES = tuple(INPUTS['r0856_table_order'])
TODAY = '2026-08-30'

# PREFLIGHT. These two are not in a default Python install, and the failure
# without them is an ImportError three frames deep rather than a sentence
# saying what to install. xlrd in particular is needed only for R0855's legacy
# .xls files, and xlrd 2.x reads .xls ONLY - openpyxl cannot stand in for it.
_missing = []
for _mod, _pkg in (('docx', 'python-docx'), ('xlrd', 'xlrd')):
    try:
        __import__(_mod)
    except ImportError:
        _missing.append(_pkg)
if _missing:
    raise SystemExit(
        'missing runtime dependencies: %s\n'
        'install them with:  python3 -m pip install -r %s\n'
        'or, if xlrd cannot be installed, convert the legacy .xls first:\n'
        '  soffice --headless --convert-to xlsx --outdir <dir> '
        'source_supplements/R0855_*.xls'
        % (', '.join(_missing),
           os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        'requirements.txt')))

import docx, xlrd
from docx.oxml.ns import qn


def sha(p):
    h = hashlib.sha256()
    with open(p, 'rb') as f:
        for b in iter(lambda: f.read(65536), b''):
            h.update(b)
    return h.hexdigest()


def cells(row):
    """The row's REAL cells. python-docx repeats a horizontally merged cell
    once per grid column it spans, which made R1040's single adjusted-HR
    column look like two identical columns - and would have written every
    estimate in eTable 7 twice."""
    out = []
    for tc in row._tr.findall(qn('w:tc')):
        out.append("".join(n.text or "" for n in tc.iter(qn('w:t'))).strip())
    return out


CI = re.compile(r'^\s*([0-9.]+)\s*\(\s*([0-9.]+)\s*[-–—]\s*([0-9.]+)\s*\)\s*$')
NDEATH = re.compile(r'^\s*([0-9]+)\s*/\s*([0-9]+)\s*$')


def parse_ci(s):
    m = CI.match(s.replace(' ', ' '))
    return (m.group(1), m.group(2), m.group(3)) if m else None


BLOCKED = []      # things this script refused to turn into a number
EFF, CNT = [], []


def eff_row(**kw):
    EFF.append(kw)


def cnt_row(**kw):
    CNT.append(kw)


# ---------------------------------------------------------------- R0856
# The record's two supplementary tables: one exposure class each, three
# stratum families per table, and two adjustment variants per stratum.
# (t4) used AFTER diagnosis. Nine strata each, and for every stratum an
# age-adjusted and a multivariable-adjusted hazard ratio.
R0856_SRC = INPUTS['sources']['R0856']
R0856_ADJ = INPUTS['adjustment_text']['R0856']
#: THE SUPPLEMENT NEVER NAMES A REFERENCE CATEGORY. Its footnote gives the
#: model covariates and nothing else; whether these hazard ratios are against
#: nonusers, against the lowest band, or against the general population is not
#: recoverable from the supplement. A reviewer has to read the main paper.
R0856_COMP = 'NOT_STATED_IN_SUPPLEMENT_REQUIRES_HUMAN_CONFIRMATION'
STRATUM = INPUTS['stratum_headers']


def open_ended(tc_row, idx):
    """True when the category cell carries a Wingdings arrow glyph.

    The top band of every stratum prints as a bare number - '7448', '9',
    '1590.00' - because the character that makes it open-ended is a Wingdings
    symbol, not text, and drops out of any text extraction. Writing '7448' as
    though it were the band would silently turn 'the highest users' into a
    point value."""
    tcs = tc_row._tr.findall(qn('w:tc'))
    return idx < len(tcs) and tcs[idx].find('.//' + qn('w:sym')) is not None


def do_R0856():
    p = os.path.join(SUPP, R0856_SRC)
    d, digest = docx.Document(p), sha(p)
    n = 0
    for ti, (t, cls) in enumerate(zip(d.tables, R0856_CLASSES)):
        supp_no = 3 + ti
        stratum = None
        for r in t.rows:
            c = cells(r)
            if len(c) < 4:
                continue
            if c[0] in STRATUM:
                stratum = STRATUM[c[0]]
                continue
            nd = NDEATH.match(c[1].replace(' ', ''))
            if not nd or stratum is None:
                continue
            label = c[0] + (' and above' if open_ended(r, 0) else '')
            if open_ended(r, 0):
                BLOCKED.append({
                    'rec_id': 'R0856', 'field': 'exposure band',
                    'where': 'supplementary table %d, %s' % (supp_no, stratum),
                    'printed': c[0],
                    'reason': 'open-ended top band; the bound symbol is a '
                              'Wingdings glyph, so whether it reads >= or > '
                              'is not recoverable from the file'})
            for qty, val in (('participants', nd.group(1)),
                             ('deaths', nd.group(2))):
                cnt_row(rec_id='R0856', quantity=qty,
                        group_role='%s users, %s %s' % (cls, stratum, label),
                        value=val, unit='persons', count_basis='raw',
                        # A COUNT BELONGS TO AN ESTIMATE. Both hazard ratios
                        # for a stratum share one denominator, so the count is
                        # anchored to the multivariable-adjusted row and says
                        # so - a count with no parent is a number floating
                        # beside a table, and two strata whose counts differ
                        # only by their parent look like one count written
                        # twice.
                        _anchor=('R0856', cls, stratum, label),
                        population_scope='stratum denominator as printed in the supplementary table. Shared by the age-adjusted and multivariable-adjusted estimates; linked to the multivariable-adjusted row as the canonical context anchor.',
                        source_page_ref='Supplementary table %d (%s), %s'
                                        % (supp_no, cls, stratum),
                        source_transcription='%s / %s (n of men / PCa deaths)'
                                             % (nd.group(1), nd.group(2)),
                        source_local_path='source_supplements/' + R0856_SRC,
                        source_sha256=digest)
            for col, variant in ((2, 'age-adjusted'), (3, 'multivariable-adjusted')):
                v = parse_ci(c[col])
                if not v:
                    BLOCKED.append({'rec_id': 'R0856', 'field': 'HR',
                                    'where': 'supp table %d, %s, %s, %s'
                                    % (supp_no, stratum, label, variant),
                                    'printed': c[col],
                                    'reason': 'cell does not parse as point (low-high)'})
                    continue
                n += 1
                eff_row(
                    rec_id='R0856', study_id=TPL['R0856']['study_id'], year=TPL['R0855']['year'],
                    study_design=TPL['R0855']['study_design'],
                    analysis_stream=TPL['R0855']['analysis_stream'],
                    synthesis_readiness=TPL['R0855']['synthesis_readiness'],
                    source_status=TPL['R0855']['source_status'],
                    source_local_path='source_supplements/' + R0856_SRC,
                    population_subgroup=TPL['R0856']['population_subgroup'],
                    exposure_class=cls,
                    exposure_definition='post-diagnostic %s use, %s: %s'
                                        % (cls, stratum, label),
                    comparator_definition=R0856_COMP,
                    comparator_type=TPL['R0856']['comparator_type'],
                    exposure_timing=TPL['R0856']['exposure_timing'],
                    outcome_name=TPL['R0856']['outcome_name'],
                    outcome_domain=TPL['R0856']['outcome_domain'],
                    outcome_definition=TPL['R0856']['outcome_definition'],
                    effect_measure='HR', effect_point=v[0],
                    effect_ci_low=v[1], effect_ci_high=v[2], ci_level='95',
                    effect_adjusted='adjusted',
                    adjustment_covariates=('age only' if variant == 'age-adjusted'
                                           else R0856_ADJ),
                    n_exposed=nd.group(1), events_exposed=nd.group(2),
                    analysis_variant='%s, %s, %s' % (stratum, label, variant),
                    data_source='supplementary table',
                    source_page_ref='Supplementary table %d (%s), %s'
                                    % (supp_no, cls, stratum),
                    source_transcription='%s, %s, %s: HR %s (95%% CI %s-%s); '
                                         'n=%s; PCa deaths=%s'
                                         % (cls, label, variant, v[0], v[1],
                                            v[2], nd.group(1), nd.group(2)),
                    source_sha256=digest,
                    _anchor_of=(('R0856', cls, stratum, label)
                                if variant == 'multivariable-adjusted' else None),
                    notes=TPL['R0856']['notes'])
    return n


# ---------------------------------------------------------------- R1040
R1040_SRC = INPUTS['sources']['R1040']
R1040_ADJ = INPUTS['adjustment_text']['R1040']


def r1040_effect(digest, **kw):
    base = dict(rec_id='R1040', study_id=TPL['R1040']['study_id'], year=TPL['R1040']['year'],
                study_design=TPL['R0855']['study_design'],
                analysis_stream=TPL['R1040']['analysis_stream'],
                synthesis_readiness=TPL['R0855']['synthesis_readiness'],
                source_status=TPL['R0855']['source_status'],
                source_local_path='source_supplements/' + R1040_SRC,
                population_subgroup=TPL['R1040']['population_subgroup'],
                exposure_class=R1040['effect_base']['exposure_class'],
                comparator_definition=R1040['effect_base']['comparator_definition'],
                comparator_type='active_comparator',
                exposure_timing=TPL['R1040']['exposure_timing'],
                outcome_name=R1040['effect_base']['outcome_name'],
                outcome_domain=R1040['effect_base']['outcome_domain'],
                outcome_definition=R1040['effect_base']['outcome_definition'],
                data_source='supplementary table', source_sha256=digest)
    base.update(kw)
    eff_row(**base)


def do_R1040():
    p = os.path.join(SUPP, R1040_SRC)
    d, digest = docx.Document(p), sha(p)
    n = 0

    # --- eTable 4: the index outcome by treatment duration and cumulative dose.
    # THE CASE COUNTS AND PERSON-YEARS ARE SUPPRESSED. Every stratum prints
    # '*' for both, which is small-cell suppression, not a missing file. The
    # hazard ratios are written; the counts are not reconstructed from the
    # rate, because a rate rounded to one decimal cannot give back a count.
    t, section = d.tables[3], None
    for r in t.rows:
        c = cells(r)
        if len(c) < 6:
            continue
        if c[0].startswith('Duration of use'):
            section = 'duration of use'
            continue
        if c[0].startswith('Cumulative Dose'):
            section = 'cumulative dose'
            continue
        if section is None or not c[0] or c[0].startswith('Abbreviations'):
            continue
        rate = parse_ci(c[3])
        if not rate:
            continue
        band = c[0]
        if c[1].strip() == '*' or c[2].strip() == '*':
            BLOCKED.append({
                'rec_id': 'R1040', 'field': 'cases and person-years',
                'where': 'eTable 4, %s, %s' % (section, band),
                'printed': '%s / %s' % (c[1], c[2]),
                'reason': 'publisher small-cell suppression; not reconstructed '
                          'from the rounded rate'})
        ref = 'Supplemental eTable 4, %s' % section
        n += 1
        r1040_effect(digest, effect_measure='incidence_rate', effect_point=rate[0],
                     effect_ci_low=rate[1], effect_ci_high=rate[2], ci_level='95',
                     effect_adjusted='crude',
                     exposure_definition='%s, %s %s' % (R1040['arms']['index'], section, band),
                     analysis_variant='%s %s, rate per 1000 person-years'
                                      % (section, band),
                     source_page_ref=ref,
                     source_transcription='%s %s: rate %s (95%% CI %s-%s) per '
                                          '1000/year' % (section, band, rate[0],
                                                         rate[1], rate[2]),
                     notes='Cases and person-years suppressed as * in the source.')
        if re.match(r'^[0-9.]+$', c[4].strip()):
            n += 1
            r1040_effect(digest, effect_measure='HR', effect_point=c[4].strip(),
                         effect_adjusted='crude',
                         exposure_definition='%s, %s %s' % (R1040['arms']['index'], section, band),
                         analysis_variant='%s %s, crude' % (section, band),
                         source_page_ref=ref,
                         source_transcription='%s %s: crude HR %s'
                                              % (section, band, c[4].strip()),
                         notes='No confidence interval printed for the crude HR.')
        adj = parse_ci(c[5])
        if adj:
            n += 1
            r1040_effect(digest, effect_measure='HR', effect_point=adj[0],
                         effect_ci_low=adj[1], effect_ci_high=adj[2], ci_level='95',
                         effect_adjusted='adjusted', adjustment_covariates=R1040_ADJ,
                         exposure_definition='%s, %s %s' % (R1040['arms']['index'], section, band),
                         analysis_variant='%s %s, adjusted' % (section, band),
                         source_page_ref=ref,
                         source_transcription='%s %s: adjusted HR %s (95%% CI %s-%s)'
                                              % (section, band, adj[0], adj[1], adj[2]))

    # --- eTable 7: the same comparison with a two-year latency window, taken
    # only for the block this review's outcome belongs to.
    t, inblock = d.tables[6], False
    for r in t.rows:
        c = cells(r)
        if len(c) < 6:
            continue
        head = c[0].strip()
        if head.endswith(LBL['r1040_block_suffixes'][0]) \
                or head == LBL['r1040_block_suffixes'][1]:
            inblock = (head == R1040['index_block_heading'])
            continue
        if not inblock or not head.startswith(tuple(R1040['arm_row_prefixes'].values())):
            continue
        arm = (R1040['arms']['reference']
               if head.startswith(R1040['arm_row_prefixes']['reference'])
               else R1040['arms']['index'])
        ref = LBL['r1040_etable7_ref']
        for qty, val, unit in (('cases', c[1].strip(), 'events'),
                               ('person_years', c[2].strip(), 'person-years')):
            if re.match(r'^[0-9,]+$', val):
                cnt_row(rec_id='R1040', quantity=qty, group_role=arm,
                        # anchored to this arm's own rate row: the cases and
                        # the person-years are what that rate was computed
                        # from, and a count with no parent is a number
                        # floating beside a table
                        _anchor=('R1040', 'etable7', arm),
                        value=val.replace(',', ''), unit=unit, count_basis='raw',
                        population_scope='male cohort, 2-year latency window',
                        source_page_ref=ref,
                        source_transcription='%s: %s %s' % (arm, val, qty),
                        source_local_path='source_supplements/' + R1040_SRC,
                        source_sha256=digest)
        rate = parse_ci(c[3])
        if rate:
            n += 1
            r1040_effect(digest, effect_measure='incidence_rate',
                         effect_point=rate[0], effect_ci_low=rate[1],
                         effect_ci_high=rate[2], ci_level='95',
                         effect_adjusted='crude', exposure_class=arm,
                         exposure_definition='%s, 2-year latency window' % arm,
                         analysis_variant='2-year latency, rate per 1000 person-years',
                         _anchor_of=('R1040', 'etable7', arm),
                         source_page_ref=ref,
                         source_transcription='%s: rate %s (95%% CI %s-%s) per 1000/year'
                                              % (arm, rate[0], rate[1], rate[2]))
        if arm == R1040['arms']['index']:
            if re.match(r'^[0-9.]+$', c[4].strip()):
                n += 1
                r1040_effect(digest, effect_measure='HR', effect_point=c[4].strip(),
                             effect_adjusted='crude',
                             exposure_definition='%s, 2-year latency window' % R1040['arms']['index'],
                             analysis_variant='2-year latency, crude',
                             source_page_ref=ref,
                             source_transcription='%s, 2-year latency: crude HR %%s' % R1040['arms']['index']
                                                  % c[4].strip(),
                             notes='No confidence interval printed for the crude HR.')
            adj = parse_ci(c[5])
            if adj:
                n += 1
                r1040_effect(digest, effect_measure='HR', effect_point=adj[0],
                             effect_ci_low=adj[1], effect_ci_high=adj[2],
                             ci_level='95', effect_adjusted='adjusted',
                             adjustment_covariates=R1040_ADJ,
                             exposure_definition='%s, 2-year latency window' % R1040['arms']['index'],
                             analysis_variant='2-year latency, adjusted',
                             source_page_ref=ref,
                             source_transcription=('%s, 2-year latency: adjusted HR %%s (95%%%% CI %%s-%%s)' % R1040['arms']['index'])
                                                  % (adj[0], adj[1], adj[2]))

    # --- eTable 1: the male denominator, and follow-up. THIS IS WHY IT
    # MATTERS: the outcome-specific covariates in eTable 1 are percentages of
    # the MALE cohort, not of the 62,109 in the header, so a reader who takes
    # the header as the denominator misreads every such row in it.
    t = d.tables[0]
    groups = [g.replace('\n', ' ').strip() for g in cells(t.rows[1])[1:]]
    for r in t.rows:
        c = cells(r)
        if len(c) < 5:
            continue
        lab = c[0].strip()
        if lab.startswith('Male, n'):
            for g, v in zip(groups, c[1:]):
                m = re.match(r'^\s*([0-9,]+)\s*\(([0-9.]+)\)\s*$', v)
                if m:
                    cnt_row(rec_id='R1040', quantity='male participants',
                            group_role=g,
                            _anchor=('R1040', 'etable1', g),
                            value=m.group(1).replace(',', ''),
                            unit='persons', count_basis='raw',
                            # THE ANCHOR IS CONTEXT, NOT AN OUTCOME. This
                            # count hangs off the group's follow-up row only
                            # because that row names the same column; saying
                            # so stops a reader taking the headcount for a
                            # follow-up-time denominator.
                            population_scope=LBL['r1040_etable1_scope'],
                            source_page_ref='Supplemental eTable 1, Male n (%)',
                            source_transcription='%s: %s (%s%%)' % (g, m.group(1), m.group(2)),
                            source_local_path='source_supplements/' + R1040_SRC,
                            source_sha256=digest)
        elif lab.startswith(('Follow-up time', 'Time at risk')):
            var = 'follow-up time' if lab.startswith('Follow-up') else 'time at risk'
            for g, v in zip(groups, c[1:]):
                m = re.match(r'^\s*([0-9.]+)\s*\(\s*([0-9.]+)\s*\)\s*$', v)
                if not m:
                    continue
                for meas, val in (('mean', m.group(1)), ('SD', m.group(2))):
                    n += 1
                    r1040_effect(digest, exposure_class=R1040['descriptive_exposure_class'],
                                 _anchor_of=(('R1040', 'etable1', g)
                                             if (var == 'follow-up time'
                                                 and meas == 'mean') else None),
                                 exposure_definition=g,
                                 comparator_definition='', comparator_type='',
                                 outcome_name=var, outcome_domain=TPL['R0855']['outcome_domain'],
                                 outcome_definition=var + ' in years',
                                 effect_measure=meas, effect_point=val,
                                 effect_adjusted='crude',
                                 analysis_variant='reported descriptive statistic',
                                 source_page_ref='Supplemental eTable 1, %s' % lab,
                                 source_transcription='%s, %s: %s (SD %s) years'
                                                      % (g, var, m.group(1), m.group(2)))
    return n


# ---------------------------------------------------------------- R0855
# Supplementary table 2: clinicopathological characteristics of the
# one treatment subgroup split by a comorbidity status. The main
# paper's Table 1 is already extracted; this is a different population and a
# different set of numbers, which is why none of it collides with what is
# there. Kept in the same shape the record's existing rows use - one row per
# median, per IQR bound - rather than a shape of my own.
R0855_SRC = INPUTS['sources']['R0855']


def do_R0855():
    p = os.path.join(SUPP, R0855_SRC)
    sh, digest = xlrd.open_workbook(p).sheet_by_index(0), sha(p)
    cell = lambda i, j: str(sh.cell_value(i, j)).strip()
    groups = {1: cell(1, 1), 3: cell(1, 3), 4: cell(1, 4)}
    n = 0
    for i in range(2, sh.nrows):
        lab = cell(i, 0)
        if not lab or not parse_ci(cell(i, 1)):
            continue
        pval = cell(i, 5)
        for col, grp in groups.items():
            v = parse_ci(cell(i, col))
            if not v:
                BLOCKED.append({'rec_id': 'R0855', 'field': 'median (IQR)',
                                'where': 'supplementary table 2, %s, %s' % (lab, grp),
                                'printed': cell(i, col),
                                'reason': 'cell does not parse as median (IQR)'})
                continue
            for meas, val in (('median', v[0]), ('IQR_low', v[1]), ('IQR_high', v[2])):
                n += 1
                eff_row(rec_id='R0855', study_id=TPL['R0855']['study_id'], year=TPL['R0855']['year'],
                        study_design=TPL['R0855']['study_design'],
                        analysis_stream=TPL['R0855']['analysis_stream'],
                        synthesis_readiness=TPL['R0855']['synthesis_readiness'],
                        source_status=TPL['R0855']['source_status'],
                        source_local_path='source_supplements/' + R0855_SRC,
                        population_subgroup=TPL['R0855']['population_subgroup'],
                        exposure_class=R1040['descriptive_exposure_class'],
                        exposure_definition=grp.replace('\n', ' '),
                        outcome_name=lab, outcome_domain=TPL['R0855']['outcome_domain'],
                        outcome_definition=lab,
                        effect_measure=meas, effect_point=val,
                        effect_adjusted='crude',
                        analysis_variant=LBL['r0855_variant'],
                        data_source='supplementary table',
                        source_page_ref=LBL['r0855_page_ref'],
                        source_transcription='%s, %s: %s (IQR %s-%s)%s'
                                             % (grp.replace('\n', ' '), lab, v[0],
                                                v[1], v[2],
                                                '; p=%s' % pval if pval else ''),
                        source_sha256=digest,
                        notes='R0855 is a qualitative include with no '
                              'class-specific antihypertensive estimate; these '
                              'are baseline descriptors, not effect estimates.')
    return n


# ---------------------------------------------------------------- assembly
HUMAN_GATE = dict(candidate_after_human_confirmation='yes', pool_eligible='no',
                  human_gate='HUMAN_DUAL_EXTRACTION_REQUIRED',
                  discrepancy_flag='no', schema_version='2',
                  evidence_capture_type='STRUCTURED_TRANSCRIPTION_NOT_VERBATIM',
                  derivation_method='reported',
                  calculation_status='REPORTED_AI_TRANSCRIBED',
                  effect_model='SEE_SOURCE_ANALYSIS_VARIANT',
                  count_scope='DESCRIPTIVE_EXPOSURE_COUNTS_NOT_TIME_VARYING_MODEL_N',
                  count_review_status='PARTIAL_REQUIRES_HUMAN_REVIEW',
                  source_quote_status='STRUCTURED_TRANSCRIPTION_NOT_VERBATIM',
                  human_confirmed='no', ai_correction_date=TODAY,
                  extraction_scope_status='PARTIAL_TEXT_TABLE_NOT_EXHAUSTIVE',
                  source_or_derivation_blocked='no')
COUNT_GATE = dict(derivation='reported',
                  evidence_status='AI_TRANSCRIBED_REQUIRES_HUMAN_REVIEW',
                  human_confirmed='no')


#: A median, an IQR bound, a mean or an SD is a DESCRIPTOR, not a reported
#: effect estimate, and the workbook's tallies split the two on this field.
#: Writing 'reported' for all of them counted 52 baseline descriptors as
#: effect estimates the first time this ran.
DESC = 'reported descriptive statistic'
DESC_MEASURES = {'median', 'IQR_low', 'IQR_high', 'mean', 'SD'}


def load(name):
    with io.open(os.path.join(BASE, name), encoding='utf-8-sig') as fh:
        rd = csv.DictReader(fh)
        return rd.fieldnames, list(rd)


def attach_count_parents(new_ef, new_ct):
    """Attach every new count to one same-record effect-context row.

    R0856 prints one denominator/event pair beside both age-adjusted and
    multivariable estimates.  The multivariable row is the canonical anchor;
    population_scope explicitly records that the count is shared.  R1040
    counts attach to the matching printed rate (eTable 7) or matching group
    follow-up mean (eTable 1).  This is a foreign-key/context link, not a claim
    that the count is an effect-model denominator.
    """
    shared_note = ('Shared by the age-adjusted and multivariable-adjusted '
                   'estimates; linked to the multivariable-adjusted row as '
                   'the canonical context anchor.')
    baseline_note = ('Baseline male headcount; the linked effect row supplies '
                     'same-column cohort context only and does not confer a '
                     'follow-up-time outcome.')

    def one(candidates, count):
        if len(candidates) != 1:
            raise AssertionError('Expected exactly one parent for %s; got %d'
                                 % (count.get('group_role'), len(candidates)))
        return candidates[0]

    for count in new_ct:
        if count['rec_id'] == 'R0856':
            context = count['group_role'].split(' users, ', 1)[1]
            norm = lambda value: re.sub(r'[\s,]+', ' ', value).strip()
            wanted = norm(context + ', multivariable-adjusted')
            parent = one([
                effect for effect in new_ef
                if effect['rec_id'] == 'R0856'
                and norm(effect['analysis_variant']) == wanted
            ], count)
            count['population_scope'] = count['population_scope'].rstrip('. ') \
                + '. ' + shared_note
        elif count['rec_id'] == 'R1040' and 'eTable 7' in count['source_page_ref']:
            parent = one([
                effect for effect in new_ef
                if effect['rec_id'] == 'R1040'
                and effect['effect_measure'] == 'incidence_rate'
                and effect['exposure_class'] == count['group_role']
                and 'eTable 7' in effect['source_page_ref']
            ], count)
        elif count['rec_id'] == 'R1040' and 'eTable 1' in count['source_page_ref']:
            parent = one([
                effect for effect in new_ef
                if effect['rec_id'] == 'R1040'
                and effect['effect_measure'] == 'mean'
                and effect['outcome_name'] == 'follow-up time'
                and effect['exposure_definition'] == count['group_role']
            ], count)
            count['population_scope'] = count['population_scope'].rstrip('. ') \
                + '. ' + baseline_note
        else:
            raise AssertionError('No count-parent rule for %s' % count)
        count['effect_row_id'] = parent['effect_row_id']
        if not count['effect_row_id'] or count['rec_id'] != parent['rec_id']:
            raise AssertionError('Invalid count parent for %s' % count)


def main():
    counts = {'R0856': do_R0856(), 'R1040': do_R1040(), 'R0855': do_R0855()}

    ef_cols, ef_rows = load('effect_extraction_text_long.csv')
    ct_cols, ct_rows = load('extraction_counts_long.csv')

    # IDs CONTINUE THE RECORD'S OWN SEQUENCE and never reuse one. A row id is
    # the key other tables point at, so a collision would silently reattach a
    # count to the wrong estimate.
    # THE ROWS FIRST, THE IDS LAST. An ordinal id can only be assigned once
    # it is known that these rows are going to be written at all; minting one
    # before the idempotency check made a rerun die on its own previous
    # output (R0856-S001 already taken) instead of reporting a clean no-op.
    new_ef = []
    for row in EFF:
        out = {c: '' for c in ef_cols}
        out.update(HUMAN_GATE)
        out.update(row)
        if out.get('_anchor_of') is None:
            out.pop('_anchor_of', None)
        if out['effect_measure'] in DESC_MEASURES:
            out['derivation_method'] = DESC
        new_ef.append(out)
    new_ct = []
    for row in CNT:
        out = {c: '' for c in ct_cols}
        out.update(COUNT_GATE)
        out.update(row)
        new_ct.append(out)

    # IDEMPOTENCY. A second --write must not append these rows again.
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    import rowkey
    # BOTH SIDES NAME THE PARENT THE SAME WAY. On file a count points at its
    # parent by ordinal; a count about to be written cannot, because the
    # ordinals do not exist yet. Give the new rows their parent's semantic key
    # from the anchor, and resolve the old rows' ordinals to the same thing.
    _parent = {}
    for out in new_ef:
        a = out.get('_anchor_of')
        if a:
            _parent[a] = str(rowkey.key_of(out, rowkey.EFFECT_KEY))
    for out in new_ef:
        out.pop('_anchor_of', None)
    for out in new_ct:
        out['parent_effect_key'] = _parent.get(out.pop('_anchor', None), '')

    # ONE PROTOCOL, SHARED. The order of a write - annotate, guard, number,
    # serialise - is where every defect that reached these tables lived, and
    # a writer that has already run never reaches it. `run_writer` is that
    # order, in one place, and `test_write_cycle` drives it through all four
    # states in CI.
    def assign_ids(label, existing, intended):
        if label == 'effects_text_long':
            nxt = {}
            for r in existing:
                m = re.match(r'^(R\d+)-T(\d+)$', r.get('effect_row_id') or '')
                if m:
                    nxt[m.group(1)] = max(nxt.get(m.group(1), 0), int(m.group(2)))
            used = {r['effect_row_id'] for r in existing}
            for out in intended:
                rid = out['rec_id']
                nxt[rid] = nxt.get(rid, 0) + 1
                eid = '%s-T%03d' % (rid, nxt[rid])
                assert eid not in used, eid
                used.add(eid)
                out['effect_row_id'] = eid
            by_key.update({str(rowkey.key_of(o, rowkey.EFFECT_KEY)):
                           o['effect_row_id'] for o in intended})
        else:
            per = {}
            for r in existing:
                m = re.match(r'^(R\d+)-S(\d+)$', r.get('count_id') or '')
                if m:
                    per[m.group(1)] = max(per.get(m.group(1), 0), int(m.group(2)))
            used_c = {r['count_id'] for r in existing}
            for out in intended:
                out['effect_row_id'] = by_key.get(out.get('parent_effect_key'), '')
                rid = out['rec_id']
                per[rid] = per.get(rid, 0) + 1
                cid = '%s-S%03d' % (rid, per[rid])
                assert cid not in used_c, cid
                used_c.add(cid)
                out['count_id'] = cid

    by_key = {}
    print('재실행 안전성 점검:')
    arch = os.path.join(BASE, 'archive', 'pre_supplement_integration_' + TODAY)
    write = '--write' in sys.argv
    verdict, written = rowkey.run_writer(
        'integrate_supplements',
        [('effects_text_long', os.path.join(BASE, 'effect_extraction_text_long.csv'),
          ef_cols, ef_rows, new_ef, rowkey.EFFECT_KEY, rowkey.EFFECT_VALUE),
         ('extraction_counts_long', os.path.join(BASE, 'extraction_counts_long.csv'),
          ct_cols, ct_rows, new_ct, rowkey.COUNT_KEY, rowkey.COUNT_VALUE)],
        bundle_paths.receipt_dir(),
        relations=[('extraction_counts_long', 'effects_text_long')],
        assignable={'effect_row_id', 'count_id', 'parent_effect_key'},
        journal_path=os.path.join(bundle_paths.receipt_dir(),
                                  'integrate_supplements_journal.json'),
        lock_path=bundle_paths.write_lock(),
        attest={'writer_code_sha256': rowkey.file_digest(
                    os.path.abspath(__file__)),
                'study_inputs_sha256': rowkey.file_digest(
                    bundle_paths.STUDY_INPUTS),
                'sources': {f: sha(os.path.join(SUPP, f))
                            for f in sorted(os.listdir(SUPP))
                            if f.startswith(('R0855', 'R0856', 'R1040'))}},
        assign_ids=assign_ids if write else None,
        archive=arch if write else None,
        dry_run=not write)
    dry = verdict != 'WRITE'
    receipt = {'date': TODAY, 'dry_run': dry, 'parsed_effect_rows': counts,
               'new_effect_rows': len(new_ef), 'new_count_rows': len(new_ct),
               'refused_values': BLOCKED, 'verdict': verdict,
               'human_fields_set_by_this_script': 'none; human_confirmed=no, '
               'pool_eligible=no, human_gate=HUMAN_DUAL_EXTRACTION_REQUIRED '
               'on every row',
               'supplement_sha256': {f: sha(os.path.join(SUPP, f))
                                     for f in sorted(os.listdir(SUPP))
                                     if f.startswith(('R0855', 'R0856', 'R1040'))}}
    if written:
        receipt['archived_to'] = 'archive/pre_supplement_integration_' + TODAY
        json.dump({'effects_text_long': [[r.get(c, '') for c in ef_cols]
                                          for r in new_ef],
                   'extraction_counts_long': [[r.get(c, '') for c in ct_cols]
                                              for r in new_ct]},
                  io.open(os.path.join(bundle_paths.receipt_dir(), 'new_rows.json'), 'w',
                          encoding='utf-8'), ensure_ascii=False)
    json.dump(receipt, io.open(os.path.join(bundle_paths.receipt_dir(), 'integration.json'), 'w',
                               encoding='utf-8'), indent=2, ensure_ascii=False)
    print(json.dumps(receipt, indent=2, ensure_ascii=False))


if __name__ == '__main__':
    main()
